from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


BASE_DIR = Path("/Users/akonkwamubagwa/Documents/Playground")
PHOTO_STRIP = BASE_DIR / "extracted_images" / "ref_page1_1_X6.png"
DOCX_OUTPUT = BASE_DIR / "Winko_Gambia_Revised_Proposal_20260316_v2.docx"
XLSX_OUTPUT = BASE_DIR / "Winko_Gambia_Revised_Tables_20260316_v2.xlsx"

RED = "D62828"
GREY = "4A4A4A"
LIGHT_GREY = "F4F4F4"
TEXT = "111111"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=TEXT, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Helvetica"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, title, rows, widths=None):
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.name = "Helvetica"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(TEXT)
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            if r_idx == 0:
                set_cell_shading(cell, GREY)
                set_cell_text(cell, value, bold=True, color=WHITE)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, LIGHT_GREY)
                set_cell_text(cell, value)
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = width
    doc.add_paragraph()


def set_run_font(run, size=10, bold=False, color=TEXT):
    run.font.name = "Helvetica"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


summary_rows = [
    ["Item", "Original Scope", "Revised Scope"],
    ["Solar PV capacity", "140 kWp", "170 kWp"],
    ["Approximate module count", "~300 modules", "~365 modules"],
    ["Usable battery capacity", "522 kWh", "622 kWh"],
    ["Hybrid inverter capacity", "125 kW", "125 kW"],
]

implementation_rows = [
    ["Phase", "Duration", "Primary Deliverables"],
    [
        "Phase I - Planning and validation",
        "2-3 weeks",
        "Final SLD, PV layout, BoM, cable routing, protection philosophy, cleaning-zone assignment, HSE and commissioning plans",
    ],
    [
        "Phase II - Procurement, logistics, and installation",
        "6-10 weeks",
        "Equipment procurement, shipping to Banjul, port handling, PV/BESS installation, cleaning-system installation, electrical integration",
    ],
    [
        "Phase III - Testing, commissioning, and handover",
        "2-3 weeks",
        "Step-load tests, grid-loss and grid-return tests, monitoring verification, training, as-builts, handover pack",
    ],
]

commercial_rows = [
    ["Category", "Cost (USD)"],
    ["Original EFSTH EPC system price (140 kWp PV + 522 kWh BESS)", "$411,900.00"],
    ["Additional scope for +30 kWp PV and +100 kWh battery", "$92,000.00"],
    ["Automated cleaning system", "$80,000.00"],
    ["Total revised project cost", "$583,900.00"],
]

payment_rows = [
    ["Time", "Phase", "Duration", "Payment trigger", "Cumulative %", "Amount (USD)", "Projected dates"],
    ["T0", "Omega", "60 days", "50% at order placement", "50.0%", "$291,950.00", "Aug 1, 2026 - Sep 30, 2026"],
    ["T1", "Sigma", "40 days", "25% before shipment", "75.0%", "$145,975.00", "Oct 1, 2026 - Nov 9, 2026"],
    ["T2", "Lambda", "30 days", "12.5% upon equipment arrival", "87.5%", "$72,987.50", "Nov 10, 2026 - Dec 9, 2026"],
    ["T3", "Kappa", "15 days", "No payment milestone", "87.5%", "$0.00", "Dec 10, 2026 - Dec 24, 2026"],
    ["T4", "Epsilon", "5 days", "Final 12.5% after reporting", "100.0%", "$72,987.50", "Dec 25, 2026 - Dec 29, 2026"],
]

bom_rows = [
    ["Hardware Category", "Cost (USD)", "Equipment Included"],
    ["PV Modules", "$69,214.29", "PV modules for approximately 365 panels, MC4 connectors, and module cabling"],
    ["Battery Energy Storage System (BESS)", "$101,283.52", "LiFePO4 battery modules, racks, BMS, and battery enclosure"],
    ["Hybrid Inverter System", "$28,000.00", "Hybrid inverter units, communication modules, and monitoring interface"],
    ["Mounting Structures & PV BOS", "$21,857.14", "Rooftop rails, clamps, structural attachments, DC combiner boxes, and PV DC cabling"],
    ["Electrical BOS & Protection", "$14,257.32", "ATS, AC cabling, disconnects, breakers, surge protection, grounding, and distribution interface"],
    ["Total hardware mapping", "$234,612.28", ""],
]

breakdown_rows = [
    ["Category Group", "Line Item", "Cost (USD)"],
    ["Hardware", "PV Modules", "$69,214.29"],
    ["Hardware", "Battery Energy Storage System (BESS)", "$101,283.52"],
    ["Hardware", "Hybrid Inverter System", "$28,000.00"],
    ["Hardware", "Mounting Structures & PV BOS", "$21,857.14"],
    ["Hardware", "Electrical BOS & Protection", "$14,257.32"],
    ["Hardware", "Total Hardware", "$234,612.28"],
    ["Non-Hardware", "Logistics & Import Handling", "$8,000.00"],
    ["Non-Hardware", "Installation & On-Site Works", "$37,100.00"],
    ["Non-Hardware", "Engineering, Travel & Commissioning", "$34,800.00"],
    ["Non-Hardware", "EPC Delivery & Performance Assurance", "$189,387.72"],
    ["Non-Hardware", "Total Non-Hardware", "$269,287.72"],
    ["Automatic Cleaning", "Automated Cleaning System", "$80,000.00"],
    ["Grand Total", "Complete Revised Project Cost", "$583,900.00"],
]


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    normal.font.size = Pt(10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("WINKO")
    set_run_font(run, size=18, bold=True, color=RED)

    p = doc.add_paragraph()
    r = p.add_run("Updated Proposal and Commercial Response")
    set_run_font(r, size=20, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("Edward Francis Small Teaching Hospital (EFSTH), Banjul, The Gambia")
    set_run_font(r, size=11, bold=True)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Solar PV, battery storage, and automated cleaning system expansion for oxygen-system reliability.")
    set_run_font(r, size=10)
    p.paragraph_format.space_after = Pt(8)

    doc.add_picture(str(PHOTO_STRIP), width=Inches(6.7))

    intro = doc.add_paragraph()
    r = intro.add_run(
        "This revised offer updates the original EFSTH EPC proposal by increasing the installed solar capacity from "
        "140 kWp to 170 kWp, increasing usable battery storage from 522 kWh to 622 kWh, and incorporating the "
        "automated PV cleaning system into the full project total."
    )
    set_run_font(r, size=10)

    add_table(doc, "Summary of Revised Scope", summary_rows)

    for heading, body, bullets in [
        (
            "1. Executive Summary",
            "The technical design basis remains tied to the oxygen plant load profile validated during the site visit on January 19, 2026: base load 15 kW, average load 25 kW, cyclic peak load 35 kW, and daily energy demand of approximately 600 kWh/day.",
            [
                "The extra 30 kWp improves daytime production margin and battery charging headroom.",
                "The extra 100 kWh improves outage-bridging capability and late-day resilience.",
                "The automated cleaning system protects real delivered energy yield under Harmattan dust and coastal salt exposure.",
                "The overall objective remains uninterrupted oxygen-system support during grid instability.",
            ],
        ),
        (
            "2. Technical Scope",
            "The revised system remains a distributed rooftop PV and battery installation across EFSTH buildings, with the battery and inverter equipment located as close as practical to the oxygen-plant point of interconnection.",
            [
                "Distributed rooftop PV arrays across the approved hospital roofs",
                "LiFePO4 battery energy storage system",
                "Hybrid inverter system with controlled solar, battery, and grid operation",
                "Dedicated oxygen-plant circuits with clear isolation and protection philosophy",
                "Remote monitoring, alarms, event logging, and cleaning-cycle reporting",
                "Automated cleaning system using rail cleaning on regular roof sections and tethered robots on irregular sections",
            ],
        ),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(heading)
        set_run_font(r, size=13, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        r = p.add_run(body)
        set_run_font(r, size=10)
        for item in bullets:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(f"• {item}")
            set_run_font(r, size=10)

    add_table(doc, "3. Implementation Plan", implementation_rows)
    add_table(doc, "4. Revised Commercial Summary", commercial_rows)
    add_table(doc, "5. Payment Timeline", payment_rows)

    p = doc.add_paragraph()
    r = p.add_run("6. Bill of Materials Included in Contract")
    set_run_font(r, size=13, bold=True)

    p = doc.add_paragraph()
    r = p.add_run(
        "Bill of Materials (BoM) and Scope Reference. The Bill of Materials attached as Appendix A forms an integral part "
        "of this Contract. The Contractor shall supply, deliver, install, test, and commission the equipment and material "
        "listed in Appendix A, together with all associated accessories, cabling, mounting, monitoring, protection devices, "
        "and incidental works required for a complete and operational system. Equivalent or higher-specification components "
        "may only be substituted with prior written approval from the Client, provided that such substitution does not "
        "reduce performance, safety, warranty coverage, or monitoring functionality."
    )
    set_run_font(r, size=10)

    add_table(doc, "Appendix A - Bill of Materials", bom_rows)
    add_table(doc, "8. Complete Project Breakdown", breakdown_rows)

    doc.save(str(DOCX_OUTPUT))


def autofit_sheet(ws):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, 14), 48)


def add_sheet(wb, name, title, rows):
    ws = wb.create_sheet(title=name)
    ws.sheet_view.showGridLines = False

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(rows[0]))
    top = ws.cell(row=1, column=1, value=title)
    top.font = Font(name="Helvetica", size=12, bold=True, color=TEXT)
    top.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 22

    header_fill = PatternFill("solid", fgColor=GREY)
    alt_fill = PatternFill("solid", fgColor=LIGHT_GREY)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for r_idx, row in enumerate(rows, start=3):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.font = Font(name="Helvetica", size=10, bold=(r_idx == 3), color=(WHITE if r_idx == 3 else TEXT))
            cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
            cell.border = border
            if r_idx == 3:
                cell.fill = header_fill
            elif r_idx % 2 == 1:
                cell.fill = alt_fill

    ws.freeze_panes = "A4"
    autofit_sheet(ws)


def build_xlsx():
    wb = Workbook()
    wb.remove(wb.active)
    add_sheet(wb, "Summary", "Summary of Revised Scope", summary_rows)
    add_sheet(wb, "Implementation", "Implementation Plan", implementation_rows)
    add_sheet(wb, "Commercial", "Revised Commercial Summary", commercial_rows)
    add_sheet(wb, "Payments", "Payment Timeline", payment_rows)
    add_sheet(wb, "BoM", "Appendix A - Bill of Materials", bom_rows)
    add_sheet(wb, "Breakdown", "Complete Project Breakdown", breakdown_rows)
    wb.save(str(XLSX_OUTPUT))


if __name__ == "__main__":
    build_docx()
    build_xlsx()
    print(DOCX_OUTPUT)
    print(XLSX_OUTPUT)
