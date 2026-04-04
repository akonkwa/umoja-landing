from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


BASE_DIR = Path("/Users/akonkwamubagwa/Documents/Playground")
PHOTO_STRIP = BASE_DIR / "extracted_images" / "liberia_ref_page1_1_X7.png"
DOCX_OUTPUT = BASE_DIR / "Winko_HopeForWomen_Liberia_Revised_Proposal_20260316.docx"
XLSX_OUTPUT = BASE_DIR / "Winko_HopeForWomen_Liberia_Revised_Tables_20260316.xlsx"

RED = "D62828"
GREY = "4A4A4A"
LIGHT_GREY = "F4F4F4"
TEXT = "111111"
WHITE = "FFFFFF"


def money(v):
    return f"${v:,.2f}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, bold=False, color=TEXT, size=12):
    run.font.name = "Helvetica"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, bold=bold)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, title, rows):
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, bold=True, size=12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            if r_idx == 0:
                set_cell_shading(cell, GREY)
                set_cell_text(cell, value, bold=True, color=WHITE)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, LIGHT_GREY)
                set_cell_text(cell, value)
    doc.add_paragraph()


def fmt_range(start, duration):
    end = start + timedelta(days=duration - 1)
    return f"{start.strftime('%b')} {start.day}, {start.year} - {end.strftime('%b')} {end.day}, {end.year}", end + timedelta(days=1)


start = date(2026, 7, 1)
t0_range, next_start = fmt_range(start, 60)
t1_range, next_start = fmt_range(next_start, 40)
t2_range, next_start = fmt_range(next_start, 30)
t3_range, next_start = fmt_range(next_start, 15)
t4_range, _ = fmt_range(next_start, 5)

original_total = 180700.00
cleaning_total = 33000.00
revised_total = original_total + cleaning_total

summary_rows = [
    ["Item", "Original Scope", "Revised Scope"],
    ["Solar PV capacity", "42 kWp", "42 kWp"],
    ["Approximate module count", "~90 modules", "~90 modules"],
    ["Usable battery capacity", "261 kWh", "261 kWh"],
    ["Automated cleaning system", "Not included", "Included"],
]

implementation_rows = [
    ["Phase", "Duration", "Primary Deliverables"],
    [
        "Phase I - Planning and validation",
        "2 weeks",
        "Final zone split, SLD confirmation, BoM, cable routing, protection philosophy, anchor points, water routing, HSE and commissioning plans",
    ],
    [
        "Phase II - Procurement, logistics, and installation",
        "6-8 weeks",
        "Equipment procurement, shipping to Monrovia, PV/BESS integration, rail fabrication, robot sourcing, installation and electrical integration",
    ],
    [
        "Phase III - Testing, commissioning, and handover",
        "2 weeks",
        "Step-load tests, cleaning-system deployment, gyro calibration, monitoring verification, training, as-builts, handover pack",
    ],
]

commercial_rows = [
    ["Category", "Cost (USD)"],
    ["Original Hope for Women EPC system price (42 kWp PV + 261 kWh BESS)", money(original_total)],
    ["Automated cleaning system", money(cleaning_total)],
    ["Total revised project cost", money(revised_total)],
]

payment_rows = [
    ["Time", "Phase", "Duration", "Payment trigger", "Cumulative %", "Amount (USD)", "Projected dates"],
    ["T0", "Omega", "60 days", "50% at order placement", "50.0%", money(revised_total * 0.5), t0_range],
    ["T1", "Sigma", "40 days", "25% before shipment", "75.0%", money(revised_total * 0.25), t1_range],
    ["T2", "Lambda", "30 days", "12.5% upon equipment arrival", "87.5%", money(revised_total * 0.125), t2_range],
    ["T3", "Kappa", "15 days", "No payment milestone", "87.5%", money(0.0), t3_range],
    ["T4", "Epsilon", "5 days", "Final 12.5% after reporting", "100.0%", money(revised_total * 0.125), t4_range],
]

bom_rows = [
    ["Hardware Category", "Cost (USD)", "Equipment Included"],
    ["PV Modules", money(18000.00), "PV modules (90 panels x 465 Wp), MC4 connectors, module cabling"],
    ["Battery Energy Storage System (BESS)", money(30000.00), "LiFePO4 battery modules, battery racks, BMS, battery cabinet/enclosure"],
    ["Hybrid Inverter System", money(12000.00), "Hybrid inverter units, communication modules, inverter monitoring interface"],
    ["Mounting Structures & PV BOS", money(8000.00), "Rooftop mounting rails, clamps, structural attachments, DC combiner boxes, PV DC cabling"],
    ["Electrical BOS & Protection", money(7000.00), "ATS, AC cabling, AC disconnects, breakers, surge protection, grounding system, distribution interface"],
    ["Total hardware mapping", money(75000.00), ""],
]

cleaning_rows = [
    ["Category", "Cost (USD)"],
    ["Zone A: Rail system (rail, carriage, brush, motor, controls)", money(8500.00)],
    ["Zone B: Tethered robot (unit, gyroscope, tether reel, dock)", money(11000.00)],
    ["Water supply (header tank, hoses, fittings, meter)", money(2000.00)],
    ["Instrumentation and monitoring integration", money(2500.00)],
    ["Installation labor (incremental to main project)", money(3000.00)],
    ["Engineering, commissioning, 6-month monitoring", money(3000.00)],
    ["Shipping (consolidated with main project)", money(3000.00)],
    ["Total pilot cleaning cost", money(33000.00)],
]

breakdown_rows = [
    ["Category Group", "Line Item", "Cost (USD)"],
    ["Hardware", "PV Modules", money(18000.00)],
    ["Hardware", "Battery Energy Storage System (BESS)", money(30000.00)],
    ["Hardware", "Hybrid Inverter System", money(12000.00)],
    ["Hardware", "Mounting Structures & PV BOS", money(8000.00)],
    ["Hardware", "Electrical BOS & Protection", money(7000.00)],
    ["Hardware", "Total Hardware", money(75000.00)],
    ["Non-Hardware", "Logistics & Import Handling", money(8000.00)],
    ["Non-Hardware", "Installation & On-Site Works", money(22300.00)],
    ["Non-Hardware", "Engineering, Travel & Commissioning", money(15400.00)],
    ["Non-Hardware", "EPC Delivery & Performance Assurance", money(60000.00)],
    ["Non-Hardware", "Total Non-Hardware", money(105700.00)],
    ["Automatic Cleaning", "Automated cleaning system", money(33000.00)],
    ["Grand Total", "Complete revised project cost", money(213700.00)],
]


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    normal.font.size = Pt(12)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("WINKO")
    set_run_font(run, bold=True, color=RED, size=18)

    p = doc.add_paragraph()
    r = p.add_run("Updated Proposal and Commercial Response")
    set_run_font(r, bold=True, size=12)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("Hope for Women Health Center, Monrovia, Liberia")
    set_run_font(r, bold=True, size=12)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Solar PV, battery storage, and automated cleaning system package for healthcare power reliability.")
    set_run_font(r, size=12)
    p.paragraph_format.space_after = Pt(8)

    doc.add_picture(str(PHOTO_STRIP), width=Inches(6.7))

    p = doc.add_paragraph()
    r = p.add_run(
        "This revised offer keeps the original Hope for Women 42 kWp solar PV and 261 kWh usable battery scope and "
        "adds the automated PV cleaning system into the full project total. The purpose is to protect real delivered "
        "generation, reduce manual roof access, and create operational data that can guide future cleaning strategy at scale."
    )
    set_run_font(r, size=12)

    add_table(doc, "Summary of Revised Scope", summary_rows)

    sections = [
        (
            "1. Executive Summary",
            "This revised package covers the original critical-load solar and battery system together with the Liberia pilot cleaning system. The clinical objective remains reliable power for imaging, sterilization, cold chain, laboratory, and other agreed essential services.",
            [
                "The base PV and battery scope remains unchanged at 42 kWp and 261 kWh usable storage.",
                "The automated cleaning system is now included in the commercial package.",
                "The cleaning system reduces safety exposure from manual roof cleaning and improves energy yield under Monrovia soiling conditions.",
                "The project still rides on the main mobilization, which keeps the cleaning add-on practical and efficient.",
            ],
        ),
        (
            "2. Technical Scope",
            "The revised scope combines the original rooftop PV and battery architecture with a two-zone automated cleaning pilot designed for the 90-module Hope for Women array.",
            [
                "Zone A uses a rail-based cleaning carriage on one half of the array.",
                "Zone B uses a tethered cleaning robot with gyroscope stabilization on the other half.",
                "The system includes water routing, cleaning controls, logging, and monitoring integration.",
                "The pilot is intended to generate measurable performance and maintenance data under real Monrovia operating conditions.",
            ],
        ),
    ]
    for heading, body, bullets in sections:
        p = doc.add_paragraph()
        r = p.add_run(heading)
        set_run_font(r, bold=True, size=12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        r = p.add_run(body)
        set_run_font(r, size=12)
        for item in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(f"• {item}")
            set_run_font(r, size=12)

    add_table(doc, "3. Implementation Plan", implementation_rows)
    add_table(doc, "4. Revised Commercial Summary", commercial_rows)
    add_table(doc, "5. Payment Timeline", payment_rows)

    p = doc.add_paragraph()
    r = p.add_run("6. Bill of Materials Included in Contract")
    set_run_font(r, bold=True, size=12)
    p = doc.add_paragraph()
    r = p.add_run(
        "Bill of Materials (BoM) and Scope Reference. The Bill of Materials attached as Appendix A forms an integral part "
        "of this Contract. The Contractor shall supply, deliver, install, test, and commission the equipment and material "
        "listed in Appendix A, together with all associated accessories, cabling, mounting, monitoring, protection devices, "
        "and incidental works required for a complete and operational system. Equivalent or higher-specification components "
        "may only be substituted with prior written approval from the Client, provided that such substitution does not "
        "reduce performance, safety, warranty coverage, or monitoring functionality."
    )
    set_run_font(r, size=12)

    add_table(doc, "Appendix A - Bill of Materials", bom_rows)
    add_table(doc, "7. Automated Cleaning System Cost", cleaning_rows)
    add_table(doc, "8. Complete Project Breakdown", breakdown_rows)
    doc.save(str(DOCX_OUTPUT))


def autofit(ws):
    for col in ws.columns:
        col_letter = get_column_letter(col[0].column)
        max_len = 0
        for cell in col:
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, 14), 48)


def add_sheet(wb, name, title, rows, currency_cols=None):
    currency_cols = currency_cols or []
    ws = wb.create_sheet(title=name)
    ws.sheet_view.showGridLines = False
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(rows[0]))
    top = ws.cell(row=1, column=1, value=title)
    top.font = Font(name="Helvetica", size=12, bold=True, color=TEXT)
    top.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 22

    header_fill = PatternFill("solid", fgColor=GREY)
    alt_fill = PatternFill("solid", fgColor=LIGHT_GREY)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for r_idx, row in enumerate(rows, start=3):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.font = Font(name="Helvetica", size=12, bold=(r_idx == 3), color=(WHITE if r_idx == 3 else TEXT))
            cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
            cell.border = border
            if r_idx == 3:
                cell.fill = header_fill
            elif r_idx % 2 == 1:
                cell.fill = alt_fill
            if c_idx - 1 in currency_cols and r_idx > 3:
                cell.number_format = '$#,##0.00'
        ws.row_dimensions[r_idx].height = 22

    ws.freeze_panes = "A4"
    autofit(ws)


def build_xlsx():
    wb = Workbook()
    wb.remove(wb.active)
    summary_x = [
        ["Item", "Original Scope", "Revised Scope"],
        ["Solar PV capacity", "42 kWp", "42 kWp"],
        ["Approximate module count", 90, 90],
        ["Usable battery capacity", "261 kWh", "261 kWh"],
        ["Automated cleaning system", "Not included", "Included"],
    ]
    implementation_x = implementation_rows
    commercial_x = [
        ["Category", "Cost (USD)"],
        ["Original Hope for Women EPC system price (42 kWp PV + 261 kWh BESS)", original_total],
        ["Automated cleaning system", cleaning_total],
        ["Total revised project cost", revised_total],
    ]
    payment_x = [
        ["Time", "Phase", "Duration", "Payment trigger", "Cumulative %", "Amount (USD)", "Projected dates"],
        ["T0", "Omega", "60 days", "50% at order placement", "50.0%", revised_total * 0.5, t0_range],
        ["T1", "Sigma", "40 days", "25% before shipment", "75.0%", revised_total * 0.25, t1_range],
        ["T2", "Lambda", "30 days", "12.5% upon equipment arrival", "87.5%", revised_total * 0.125, t2_range],
        ["T3", "Kappa", "15 days", "No payment milestone", "87.5%", 0.0, t3_range],
        ["T4", "Epsilon", "5 days", "Final 12.5% after reporting", "100.0%", revised_total * 0.125, t4_range],
    ]
    bom_x = [
        ["Hardware Category", "Cost (USD)", "Equipment Included"],
        ["PV Modules", 18000.00, "PV modules (90 panels x 465 Wp), MC4 connectors, module cabling"],
        ["Battery Energy Storage System (BESS)", 30000.00, "LiFePO4 battery modules, battery racks, BMS, battery cabinet/enclosure"],
        ["Hybrid Inverter System", 12000.00, "Hybrid inverter units, communication modules, inverter monitoring interface"],
        ["Mounting Structures & PV BOS", 8000.00, "Rooftop mounting rails, clamps, structural attachments, DC combiner boxes, PV DC cabling"],
        ["Electrical BOS & Protection", 7000.00, "ATS, AC cabling, AC disconnects, breakers, surge protection, grounding system, distribution interface"],
        ["Total hardware mapping", 75000.00, ""],
    ]
    cleaning_x = [
        ["Category", "Cost (USD)"],
        ["Zone A: Rail system (rail, carriage, brush, motor, controls)", 8500.00],
        ["Zone B: Tethered robot (unit, gyroscope, tether reel, dock)", 11000.00],
        ["Water supply (header tank, hoses, fittings, meter)", 2000.00],
        ["Instrumentation and monitoring integration", 2500.00],
        ["Installation labor (incremental to main project)", 3000.00],
        ["Engineering, commissioning, 6-month monitoring", 3000.00],
        ["Shipping (consolidated with main project)", 3000.00],
        ["Total pilot cleaning cost", 33000.00],
    ]
    breakdown_x = [
        ["Category Group", "Line Item", "Cost (USD)"],
        ["Hardware", "PV Modules", 18000.00],
        ["Hardware", "Battery Energy Storage System (BESS)", 30000.00],
        ["Hardware", "Hybrid Inverter System", 12000.00],
        ["Hardware", "Mounting Structures & PV BOS", 8000.00],
        ["Hardware", "Electrical BOS & Protection", 7000.00],
        ["Hardware", "Total Hardware", 75000.00],
        ["Non-Hardware", "Logistics & Import Handling", 8000.00],
        ["Non-Hardware", "Installation & On-Site Works", 22300.00],
        ["Non-Hardware", "Engineering, Travel & Commissioning", 15400.00],
        ["Non-Hardware", "EPC Delivery & Performance Assurance", 60000.00],
        ["Non-Hardware", "Total Non-Hardware", 105700.00],
        ["Automatic Cleaning", "Automated cleaning system", 33000.00],
        ["Grand Total", "Complete revised project cost", 213700.00],
    ]

    add_sheet(wb, "Summary", "Summary of Revised Scope", summary_x)
    add_sheet(wb, "Implementation", "Implementation Plan", implementation_x)
    add_sheet(wb, "Commercial", "Revised Commercial Summary", commercial_x, currency_cols=[1])
    add_sheet(wb, "Payments", "Payment Timeline", payment_x, currency_cols=[5])
    add_sheet(wb, "BoM", "Appendix A - Bill of Materials", bom_x, currency_cols=[1])
    add_sheet(wb, "Cleaning", "Automated Cleaning System Cost", cleaning_x, currency_cols=[1])
    add_sheet(wb, "Breakdown", "Complete Project Breakdown", breakdown_x, currency_cols=[2])
    wb.save(str(XLSX_OUTPUT))


if __name__ == "__main__":
    build_docx()
    build_xlsx()
    print(DOCX_OUTPUT)
    print(XLSX_OUTPUT)
