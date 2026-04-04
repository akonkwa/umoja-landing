from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PLAYGROUND = Path("/Users/akonkwamubagwa/Documents/Playground")
DOCUMENTS = Path("/Users/akonkwamubagwa/Documents")
PHOTO_STRIP = PLAYGROUND / "extracted_images" / "ref_page1_1_X6.png"
EXPANDED_BOM = DOCUMENTS / "Winko Expanded BoM 20260217.xlsx"

DOCX_OUTPUT = DOCUMENTS / "Winko IGPC SCCM Gambia EFSTH Banjul Revised Proposal 20260317 Expanded BoM.docx"
PDF_OUTPUT = DOCUMENTS / "Winko IGPC SCCM Gambia EFSTH Banjul Revised Proposal 20260317 Expanded BoM.pdf"
XLSX_OUTPUT = DOCUMENTS / "Winko IGPC SCCM Gambia EFSTH Banjul Revised Proposal 20260317 Expanded BoM Tables.xlsx"

RED = "D62828"
GREY = "4A4A4A"
LIGHT_GREY = "F4F4F4"
TEXT = "111111"
WHITE = "FFFFFF"

RL_RED = colors.HexColor("#D62828")
RL_GREY = colors.HexColor("#4A4A4A")
RL_LIGHT_GREY = colors.HexColor("#F4F4F4")
RL_TEXT = colors.HexColor("#111111")
RL_MUTED = colors.HexColor("#5F6368")
RL_LINE = colors.HexColor("#D9D9D9")


def money(v):
    return f"${v:,.2f}"


def load_expanded_rows():
    wb = load_workbook(EXPANDED_BOM, data_only=True)
    ws = wb["Expanded_BoM"]
    rows = []
    grand_total = None
    for row in ws.iter_rows(values_only=True):
        if not any(v is not None for v in row):
            continue
        _, category, subcomponent, description, cost = row
        if category == "Category":
            continue
        if category is None and cost is not None:
            grand_total = round(float(cost), 2)
            continue
        rows.append([
            "" if category is None else category,
            "" if subcomponent is None else subcomponent,
            "" if description is None else description,
            money(round(float(cost), 2)) if cost is not None else "",
        ])
    return rows, grand_total


expanded_rows, expanded_total = load_expanded_rows()

summary_rows = [
    ["Item", "Original Scope", "Revised Scope"],
    ["Solar PV capacity", "140 kWp", "170 kWp"],
    ["Approximate module count", "~300 modules", "~365 modules"],
    ["Usable battery capacity", "522 kWh", "622 kWh"],
    ["Hybrid inverter capacity", "125 kW", "125 kW"],
]

implementation_rows = [
    ["Phase", "Duration", "Primary Deliverables"],
    ["Phase I - Planning and validation", "2-3 weeks", "Final SLD, PV layout, BoM, cable routing, protection philosophy, cleaning-zone assignment, HSE and commissioning plans"],
    ["Phase II - Procurement, logistics, and installation", "6-10 weeks", "Equipment procurement, shipping to Banjul, port handling, PV/BESS installation, cleaning-system installation, electrical integration"],
    ["Phase III - Testing, commissioning, and handover", "2-3 weeks", "Step-load tests, grid-loss and grid-return tests, monitoring verification, training, as-builts, handover pack"],
]

commercial_rows = [
    ["Category", "Cost (USD)"],
    ["Original EFSTH EPC system price (140 kWp PV + 522 kWh BESS)", money(411900.00)],
    ["Additional scope for +30 kWp PV and +100 kWh battery", money(92000.00)],
    ["Automated cleaning system", money(80000.00)],
    ["Total revised project cost", money(583900.00)],
]

payment_rows = [
    ["Time", "Phase", "Duration", "Payment trigger", "Cumulative %", "Amount (USD)", "Projected dates"],
    ["T0", "Omega", "60 days", "50% at order placement", "50.0%", money(291950.00), "Aug 1, 2026 - Sep 30, 2026"],
    ["T1", "Sigma", "40 days", "25% before shipment", "75.0%", money(145975.00), "Oct 1, 2026 - Nov 9, 2026"],
    ["T2", "Lambda", "30 days", "12.5% upon equipment arrival", "87.5%", money(72987.50), "Nov 10, 2026 - Dec 9, 2026"],
    ["T3", "Kappa", "15 days", "No payment milestone", "87.5%", money(0.00), "Dec 10, 2026 - Dec 24, 2026"],
    ["T4", "Epsilon", "5 days", "Final 12.5% after reporting", "100.0%", money(72987.50), "Dec 25, 2026 - Dec 29, 2026"],
]

bom_rows = [
    ["Hardware Category", "Cost (USD)", "Equipment Included"],
    ["PV Modules", money(69214.29), "PV modules for approximately 365 panels, MC4 connectors, and module cabling"],
    ["Battery Energy Storage System (BESS)", money(101283.52), "LiFePO4 battery modules, battery racks, BMS, and battery enclosure"],
    ["Hybrid Inverter System", money(28000.00), "Hybrid inverter units, communication modules, and monitoring interface"],
    ["Mounting Structures & PV BOS", money(21857.14), "Rooftop rails, clamps, structural attachments, DC combiner boxes, and PV DC cabling"],
    ["Electrical BOS & Protection", money(14257.32), "ATS, AC cabling, disconnects, breakers, surge protection, grounding, and distribution interface"],
    ["Total hardware mapping", money(234612.28), ""],
]

breakdown_rows = [
    ["Category Group", "Line Item", "Cost (USD)"],
    ["Hardware", "PV Modules", money(69214.29)],
    ["Hardware", "Battery Energy Storage System (BESS)", money(101283.52)],
    ["Hardware", "Hybrid Inverter System", money(28000.00)],
    ["Hardware", "Mounting Structures & PV BOS", money(21857.14)],
    ["Hardware", "Electrical BOS & Protection", money(14257.32)],
    ["Hardware", "Total Hardware", money(234612.28)],
    ["Non-Hardware", "Logistics & Import Handling", money(8000.00)],
    ["Non-Hardware", "Installation & On-Site Works", money(37100.00)],
    ["Non-Hardware", "Engineering, Travel & Commissioning", money(34800.00)],
    ["Non-Hardware", "EPC Delivery & Performance Assurance", money(189387.72)],
    ["Non-Hardware", "Total Non-Hardware", money(269287.72)],
    ["Automatic Cleaning", "Automated Cleaning System", money(80000.00)],
    ["Grand Total", "Complete Revised Project Cost", money(583900.00)],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=10, bold=False, color=TEXT):
    run.font.name = "Helvetica"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False, color=TEXT, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(doc, title, rows):
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            if r_idx == 0:
                set_cell_shading(cell, GREY)
                set_cell_text(cell, value, bold=True, color=WHITE)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, LIGHT_GREY)
                set_cell_text(cell, value)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Helvetica")
    normal.font.size = Pt(10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("WINKO")
    set_run_font(run, size=18, bold=True, color=RED)

    p = doc.add_paragraph()
    r = p.add_run("Updated Proposal and Commercial Response")
    set_run_font(r, size=20, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("Edward Francis Small Teaching Hospital (EFSTH), Banjul, The Gambia")
    set_run_font(r, size=11, bold=True)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Solar PV, battery storage, and automated cleaning system expansion for oxygen-system reliability.")
    set_run_font(r, size=10)
    p.paragraph_format.space_after = Pt(8)

    doc.add_picture(str(PHOTO_STRIP), width=Inches(6.7))

    p = doc.add_paragraph()
    r = p.add_run(
        "This revised offer updates the original EFSTH EPC proposal by increasing the installed solar capacity from 140 kWp "
        "to 170 kWp, increasing usable battery storage from 522 kWh to 622 kWh, incorporating the automated PV cleaning "
        "system into the full project total, and attaching the expanded bill of materials as a formal appendix."
    )
    set_run_font(r, size=10)

    add_docx_table(doc, "Summary of Revised Scope", summary_rows)

    sections = [
        (
            "1. Executive Summary",
            "The technical design basis remains tied to the oxygen plant load profile validated during the site visit on January 19, 2026: base load 15 kW, average load 25 kW, cyclic peak load 35 kW, and daily energy demand of approximately 600 kWh/day.",
            [
                "The extra 30 kWp improves daytime production margin and battery charging headroom.",
                "The extra 100 kWh improves outage-bridging capability and late-day resilience.",
                "The automated cleaning system protects real delivered energy yield under Harmattan dust and coastal salt exposure.",
                "The expanded BoM now shows the project cost logic in more detail while keeping the scope aligned to system delivery rather than supplier-level part resale.",
            ],
        ),
        (
            "2. Technical Scope",
            "The revised system remains a distributed rooftop PV and battery installation across EFSTH buildings, with the battery and inverter equipment located as close as practical to the oxygen-plant point of interconnection.",
            [
                "Distributed rooftop PV arrays across the approved hospital roofs",
                "LiFePO4 battery energy storage system",
                "Hybrid inverter system with controlled solar, battery, and grid operation",
                "Dedicated oxygen-plant circuits with clear isolation and protection philosophy",
                "Remote monitoring, alarms, event logging, and cleaning-cycle reporting",
                "Automated cleaning system using rail cleaning on regular roof sections and tethered robots on irregular sections",
            ],
        ),
    ]
    for heading, body, bullets in sections:
        p = doc.add_paragraph()
        r = p.add_run(heading)
        set_run_font(r, size=13, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        r = p.add_run(body)
        set_run_font(r, size=10)
        for item in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(f"• {item}")
            set_run_font(r, size=10)

    add_docx_table(doc, "3. Implementation Plan", implementation_rows)
    add_docx_table(doc, "4. Revised Commercial Summary", commercial_rows)
    add_docx_table(doc, "5. Payment Timeline", payment_rows)

    p = doc.add_paragraph()
    r = p.add_run("6. Bill of Materials Included in Contract")
    set_run_font(r, size=13, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(
        "Bill of Materials (BoM) and Scope Reference. The Bill of Materials attached as Appendix A and the Expanded BoM "
        "attached as Appendix B form integral parts of this Contract. The Contractor shall supply, deliver, install, test, "
        "and commission the equipment and material listed in those appendices, together with all associated accessories, "
        "cabling, mounting, monitoring, protection devices, and incidental works required for a complete and operational system."
    )
    set_run_font(r, size=10)

    add_docx_table(doc, "Appendix A - Bill of Materials", bom_rows)
    add_docx_table(doc, "8. Complete Project Breakdown", breakdown_rows)
    add_docx_table(doc, "Appendix B - Expanded Bill of Materials", [["Category", "Subcomponent", "Description", "Cost (USD)"]] + expanded_rows + [["", "", "Expanded BoM total", money(expanded_total or 583900.00)]])

    doc.save(str(DOCX_OUTPUT))


styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="BodyC", fontName="Helvetica", fontSize=10, leading=14, textColor=RL_TEXT, spaceAfter=6))
styles.add(ParagraphStyle(name="SmallC", fontName="Helvetica", fontSize=8.3, leading=10.5, textColor=RL_TEXT, spaceAfter=4))
styles.add(ParagraphStyle(name="TitleC", fontName="Helvetica-Bold", fontSize=20, leading=25, textColor=RL_TEXT, spaceAfter=10))
styles.add(ParagraphStyle(name="SectionC", fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=RL_TEXT, spaceAfter=8, spaceBefore=8))
styles.add(ParagraphStyle(name="SubsectionC", fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=RL_TEXT, spaceAfter=6))


def header(canvas, doc):
    width, height = A4
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 18)
    canvas.setFillColor(RL_RED)
    canvas.drawRightString(width - doc.rightMargin, height - 36, "WINKO")
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(RL_MUTED)
    canvas.drawString(doc.leftMargin, 22, "Edward Francis Small Teaching Hospital (EFSTH) | Revised Proposal")
    canvas.drawRightString(width - doc.rightMargin, 22, str(canvas.getPageNumber()))
    canvas.restoreState()


def p(text, style="BodyC"):
    return Paragraph(text, styles[style])


def bullet(text):
    return Paragraph(f"• {text}", styles["BodyC"])


def make_pdf_table(rows, col_widths, small=False):
    style = styles["SmallC"] if small else styles["BodyC"]
    wrapped = [[Paragraph(str(cell), style) for cell in row] for row in rows]
    table = Table(wrapped, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), RL_GREY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, RL_LINE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, RL_LIGHT_GREY]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf():
    story = []
    story.append(Spacer(1, 0.2 * inch))
    story.append(p("Updated Proposal and Commercial Response", "TitleC"))
    story.append(p("Edward Francis Small Teaching Hospital (EFSTH), Banjul, The Gambia", "SubsectionC"))
    story.append(p("Solar PV, battery storage, and automated cleaning system expansion for oxygen-system reliability."))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Image(str(PHOTO_STRIP), width=6.9 * inch, height=1.75 * inch))
    story.append(Spacer(1, 0.14 * inch))
    story.append(p("This revised offer updates the original EFSTH EPC proposal by increasing the installed solar capacity from <b>140 kWp</b> to <b>170 kWp</b>, increasing usable battery storage from <b>522 kWh</b> to <b>622 kWh</b>, incorporating the automated PV cleaning system into the full project total, and attaching the expanded bill of materials as a formal appendix."))
    story.append(make_pdf_table(summary_rows, [2.5 * inch, 2.0 * inch, 2.0 * inch]))
    story.append(p("1. Executive Summary", "SectionC"))
    story.append(p("The technical design basis remains tied to the oxygen plant load profile validated during the site visit on <b>January 19, 2026</b>: base load <b>15 kW</b>, average load <b>25 kW</b>, cyclic peak load <b>35 kW</b>, and daily energy demand of approximately <b>600 kWh/day</b>."))
    for item in [
        "The extra 30 kWp improves daytime production margin and battery charging headroom.",
        "The extra 100 kWh improves outage-bridging capability and late-day resilience.",
        "The automated cleaning system protects real delivered energy yield under Harmattan dust and coastal salt exposure.",
        "The expanded BoM now shows the project cost logic in more detail while keeping the scope aligned to system delivery rather than supplier-level part resale.",
    ]:
        story.append(bullet(item))
    story.append(p("2. Technical Scope", "SectionC"))
    story.append(p("The revised system remains a distributed rooftop PV and battery installation across EFSTH buildings, with the battery and inverter equipment located as close as practical to the oxygen-plant point of interconnection."))
    for item in [
        "Distributed rooftop PV arrays across the approved hospital roofs",
        "LiFePO4 battery energy storage system",
        "Hybrid inverter system with controlled solar, battery, and grid operation",
        "Dedicated oxygen-plant circuits with clear isolation and protection philosophy",
        "Remote monitoring, alarms, event logging, and cleaning-cycle reporting",
        "Automated cleaning system using rail cleaning on regular roof sections and tethered robots on irregular sections",
    ]:
        story.append(bullet(item))
    story.append(p("3. Implementation Plan", "SectionC"))
    story.append(make_pdf_table(implementation_rows, [2.0 * inch, 1.0 * inch, 3.7 * inch], small=True))
    story.append(p("4. Revised Commercial Summary", "SectionC"))
    story.append(make_pdf_table(commercial_rows, [4.8 * inch, 1.9 * inch]))
    story.append(p("5. Payment Timeline", "SectionC"))
    story.append(make_pdf_table(payment_rows, [0.45 * inch, 0.7 * inch, 0.7 * inch, 1.8 * inch, 0.7 * inch, 1.0 * inch, 1.55 * inch], small=True))
    story.append(PageBreak())
    story.append(p("6. Bill of Materials Included in Contract", "SectionC"))
    story.append(p("Bill of Materials (BoM) and Scope Reference. The Bill of Materials attached as Appendix A and the Expanded BoM attached as Appendix B form integral parts of this Contract. The Contractor shall supply, deliver, install, test, and commission the equipment and material listed in those appendices, together with all associated accessories, cabling, mounting, monitoring, protection devices, and incidental works required for a complete and operational system."))
    story.append(p("Appendix A - Bill of Materials", "SectionC"))
    story.append(make_pdf_table(bom_rows, [2.2 * inch, 1.0 * inch, 3.45 * inch], small=True))
    story.append(p("8. Complete Project Breakdown", "SectionC"))
    story.append(make_pdf_table(breakdown_rows, [1.35 * inch, 4.1 * inch, 1.25 * inch], small=True))
    story.append(PageBreak())
    story.append(p("Appendix B - Expanded Bill of Materials", "SectionC"))
    expanded_table_rows = [["Category", "Subcomponent", "Description", "Cost (USD)"]] + expanded_rows + [["", "", "Expanded BoM total", money(expanded_total or 583900.00)]]
    story.append(make_pdf_table(expanded_table_rows, [1.5 * inch, 1.7 * inch, 2.9 * inch, 0.9 * inch], small=True))

    doc = SimpleDocTemplate(str(PDF_OUTPUT), pagesize=A4, leftMargin=42, rightMargin=42, topMargin=56, bottomMargin=34)
    doc.build(story, onFirstPage=header, onLaterPages=header)


def autofit_sheet(ws):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, 14), 52)


def add_sheet(wb, name, title, rows, currency_cols=None):
    currency_cols = currency_cols or []
    ws = wb.create_sheet(title=name)
    ws.sheet_view.showGridLines = False
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(rows[0]))
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(name="Helvetica", size=12, bold=True, color=TEXT)
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 22

    header_fill = PatternFill("solid", fgColor=GREY)
    alt_fill = PatternFill("solid", fgColor=LIGHT_GREY)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for r_idx, row in enumerate(rows, start=3):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.font = Font(
                name="Helvetica",
                size=10,
                bold=(r_idx == 3),
                color=(WHITE if r_idx == 3 else TEXT),
            )
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = border
            if r_idx == 3:
                cell.fill = header_fill
            elif r_idx % 2 == 1:
                cell.fill = alt_fill
            if r_idx > 3 and (c_idx - 1) in currency_cols and isinstance(value, (int, float)):
                cell.number_format = '$#,##0.00'
        ws.row_dimensions[r_idx].height = 22

    ws.freeze_panes = "A4"
    autofit_sheet(ws)


def build_xlsx():
    wb = Workbook()
    wb.remove(wb.active)

    summary_x = [
        ["Item", "Original Scope", "Revised Scope"],
        ["Solar PV capacity", "140 kWp", "170 kWp"],
        ["Approximate module count", "~300 modules", "~365 modules"],
        ["Usable battery capacity", "522 kWh", "622 kWh"],
        ["Hybrid inverter capacity", "125 kW", "125 kW"],
    ]
    commercial_x = [
        ["Category", "Cost (USD)"],
        ["Original EFSTH EPC system price (140 kWp PV + 522 kWh BESS)", 411900.00],
        ["Additional scope for +30 kWp PV and +100 kWh battery", 92000.00],
        ["Automated cleaning system", 80000.00],
        ["Total revised project cost", 583900.00],
    ]
    payment_x = [
        ["Time", "Phase", "Duration", "Payment trigger", "Cumulative %", "Amount (USD)", "Projected dates"],
        ["T0", "Omega", "60 days", "50% at order placement", "50.0%", 291950.00, "Aug 1, 2026 - Sep 30, 2026"],
        ["T1", "Sigma", "40 days", "25% before shipment", "75.0%", 145975.00, "Oct 1, 2026 - Nov 9, 2026"],
        ["T2", "Lambda", "30 days", "12.5% upon equipment arrival", "87.5%", 72987.50, "Nov 10, 2026 - Dec 9, 2026"],
        ["T3", "Kappa", "15 days", "No payment milestone", "87.5%", 0.00, "Dec 10, 2026 - Dec 24, 2026"],
        ["T4", "Epsilon", "5 days", "Final 12.5% after reporting", "100.0%", 72987.50, "Dec 25, 2026 - Dec 29, 2026"],
    ]
    bom_x = [
        ["Hardware Category", "Cost (USD)", "Equipment Included"],
        ["PV Modules", 69214.29, "PV modules for approximately 365 panels, MC4 connectors, and module cabling"],
        ["Battery Energy Storage System (BESS)", 101283.52, "LiFePO4 battery modules, battery racks, BMS, and battery enclosure"],
        ["Hybrid Inverter System", 28000.00, "Hybrid inverter units, communication modules, and monitoring interface"],
        ["Mounting Structures & PV BOS", 21857.14, "Rooftop rails, clamps, structural attachments, DC combiner boxes, and PV DC cabling"],
        ["Electrical BOS & Protection", 14257.32, "ATS, AC cabling, disconnects, breakers, surge protection, grounding, and distribution interface"],
        ["Total hardware mapping", 234612.28, ""],
    ]
    breakdown_x = [
        ["Category Group", "Line Item", "Cost (USD)"],
        ["Hardware", "PV Modules", 69214.29],
        ["Hardware", "Battery Energy Storage System (BESS)", 101283.52],
        ["Hardware", "Hybrid Inverter System", 28000.00],
        ["Hardware", "Mounting Structures & PV BOS", 21857.14],
        ["Hardware", "Electrical BOS & Protection", 14257.32],
        ["Hardware", "Total Hardware", 234612.28],
        ["Non-Hardware", "Logistics & Import Handling", 8000.00],
        ["Non-Hardware", "Installation & On-Site Works", 37100.00],
        ["Non-Hardware", "Engineering, Travel & Commissioning", 34800.00],
        ["Non-Hardware", "EPC Delivery & Performance Assurance", 189387.72],
        ["Non-Hardware", "Total Non-Hardware", 269287.72],
        ["Automatic Cleaning", "Automated Cleaning System", 80000.00],
        ["Grand Total", "Complete Revised Project Cost", 583900.00],
    ]
    expanded_x = [["Category", "Subcomponent", "Description", "Cost (USD)"]]
    for row in expanded_rows:
        category, subcomponent, description, cost = row
        numeric_cost = float(cost.replace("$", "").replace(",", "")) if cost else None
        expanded_x.append([category, subcomponent, description, numeric_cost])
    expanded_x.append(["", "", "Expanded BoM total", expanded_total or 583900.00])

    add_sheet(wb, "Summary", "Summary of Revised Scope", summary_x)
    add_sheet(wb, "Implementation", "Implementation Plan", implementation_rows)
    add_sheet(wb, "Commercial", "Revised Commercial Summary", commercial_x, currency_cols=[1])
    add_sheet(wb, "Payments", "Payment Timeline", payment_x, currency_cols=[5])
    add_sheet(wb, "BoM", "Appendix A - Bill of Materials", bom_x, currency_cols=[1])
    add_sheet(wb, "Breakdown", "Complete Project Breakdown", breakdown_x, currency_cols=[2])
    add_sheet(wb, "Expanded BoM", "Appendix B - Expanded Bill of Materials", expanded_x, currency_cols=[3])
    wb.save(str(XLSX_OUTPUT))


if __name__ == "__main__":
    build_docx()
    build_pdf()
    build_xlsx()
    print(DOCX_OUTPUT)
    print(PDF_OUTPUT)
    print(XLSX_OUTPUT)
